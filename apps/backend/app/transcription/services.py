"""Transcription services."""

import asyncio
from datetime import datetime
from typing import Optional, Tuple, List, Dict, Any
from flask import current_app
from app.extensions import db
from app.transcription.models import Transcription, TranscriptionSegment
from app.transcription.repositories import TranscriptionRepository
from app.transcription.ai_client import AIServiceClient
from app.audio.models import AudioFile
from app.utils.logging_middleware import log_external_service_call
from app.cache.redis_service import get_transcription_cache
import logging

logger = logging.getLogger(__name__)


class TranscriptionService:
    """Service for transcription operations."""
    
    def __init__(self):
        self.repository = TranscriptionRepository()
        self.ai_client = AIServiceClient()
        # Note: cache is accessed via method to ensure fresh instances
    
    def create_transcription_with_metadata(self, audio_file_id: int, user_id: int,
                                          title: str, description: str = '',
                                          language: str = 'el',
                                          ai_model: str = 'whisper-large-v3') -> Transcription:
        """Create a new transcription job with metadata."""
        logger.info(f"Creating transcription job for audio file {audio_file_id}")
        
        audio_file = AudioFile.query.get(audio_file_id)
        if not audio_file:
            logger.error(f"Audio file not found | audio_file_id={audio_file_id}")
            raise ValueError("Audio file not found")
        
        logger.info(f"Audio file loaded: {audio_file.original_filename}")
        
        transcription = Transcription(
            audio_file_id=audio_file_id,
            user_id=user_id,
            title=title,
            description=description,
            language=language,
            model_used=ai_model,
            status='pending',
            duration_seconds=audio_file.duration_seconds
        )
        
        db.session.add(transcription)
        db.session.commit()
        
        # Invalidate analytics cache when new transcription is created
        try:
            from app.analytics.services import AnalyticsService
            analytics_service = AnalyticsService()
            # Invalidate user-specific cache 
            analytics_service._invalidate_cache_pattern(f'user_stats:user_id={user_id}')
            logger.info(f"Analytics cache invalidated for new transcription by user {user_id}")
        except Exception as cache_error:
            logger.warning(f"Failed to invalidate analytics cache: {str(cache_error)}")
        
        logger.info(f"Transcription job created with metadata | transcription_id={transcription.id} | status=pending")
        
        self._process_transcription_async(transcription.id)
        
        return transcription

    def create_transcription(self, audio_file_id: int, user_id: int, 
                           language: str = 'el') -> Transcription:
        """Create a new transcription job."""
        logger.info(f"Creating transcription job | audio_file_id={audio_file_id} | user_id={user_id} | language={language}")
        
        audio_file = AudioFile.query.get(audio_file_id)
        if not audio_file:
            logger.error(f"Audio file not found | audio_file_id={audio_file_id}")
            raise ValueError("Audio file not found")
        
        logger.info(f"Audio file loaded: {audio_file.original_filename}")
        
        transcription = Transcription(
            audio_file_id=audio_file_id,
            user_id=user_id,
            language=language,
            status='pending',
            duration_seconds=audio_file.duration_seconds
        )
        
        db.session.add(transcription)
        db.session.commit()
        
        # Invalidate analytics cache when new transcription is created
        try:
            from app.analytics.services import AnalyticsService
            analytics_service = AnalyticsService()
            # Invalidate user-specific cache 
            analytics_service._invalidate_cache_pattern(f'user_stats:user_id={user_id}')
            logger.info(f"Analytics cache invalidated for new transcription by user {user_id}")
        except Exception as cache_error:
            logger.warning(f"Failed to invalidate analytics cache: {str(cache_error)}")
        
        logger.info(f"Transcription job created | transcription_id={transcription.id} | status=pending")
        
        self._process_transcription_async(transcription.id)
        
        return transcription
    
    def _process_transcription_async(self, transcription_id: int):
        """Process transcription asynchronously."""
        import threading
        from flask import current_app
        
        app = current_app._get_current_object()
        
        def process():
            with app.app_context():
                self._process_transcription(transcription_id, app)
        
        thread = threading.Thread(target=process)
        thread.daemon = True
        thread.start()
    
    def _process_transcription(self, transcription_id: int, app=None):
        """Process the transcription using AI service with real-time progress updates."""
        transcription = Transcription.query.get(transcription_id)
        if not transcription:
            return
        
        from app.websocket.manager import get_progress_manager
        progress_manager = get_progress_manager()
        
        from app.extensions import socketio
        
        def on_progress(stage: str, percentage: int, message: str):
            """Progress callback for real-time updates via WebSocket."""
            if progress_manager and socketio:
                try:
                    progress_data = {
                        'stage': stage,
                        'percentage': percentage,
                        'message': message,
                        'transcription_id': str(transcription_id)
                    }
                    if app:
                        with app.app_context():
                            progress_manager.broadcast_progress(str(transcription_id), progress_data)
                    else:
                        progress_manager.broadcast_progress(str(transcription_id), progress_data)
                    logger.info(f"Progress update: {transcription_id} | {stage} | {percentage}%")
                except Exception as e:
                    logger.error(f"Progress broadcast error | transcription_id={transcription_id} | error={str(e)}")
        
        try:
            on_progress('initializing', 5, 'Ξεκινά η επεξεργασία...')
            
            transcription.status = 'processing'
            transcription.started_at = datetime.utcnow()
            db.session.commit()
            
            on_progress('preprocessing', 15, 'Προετοιμασία αρχείου ήχου...')
            
            audio_file = transcription.audio_file
            template = None
            
            on_progress('ai_processing', 25, 'Φόρτωση μοντέλου AI...')
            
            model = 'whisper'
            if transcription.model_used:
                if 'whisper' in transcription.model_used.lower():
                    model = 'whisper'
                elif 'wav2vec' in transcription.model_used.lower():
                    model = 'wav2vec2'
                elif transcription.model_used in ['both', 'compare']:
                    model = 'both'
            
            if model == 'both':
                on_progress('ai_processing', 35, 'Επεξεργασία με και τα δύο μοντέλα...')
                
                if app:
                    with app.app_context():
                        result = self.ai_client.transcribe_both_models(
                            audio_file_path=audio_file.file_path,
                            language=transcription.language,
                            original_filename=audio_file.original_filename
                        )
                else:
                    result = self.ai_client.transcribe_both_models(
                        audio_file_path=audio_file.file_path,
                        language=transcription.language,
                        original_filename=audio_file.original_filename
                    )
                
                # Store results for both models
                transcription.whisper_text = result.get('whisper_text', '')
                transcription.whisper_confidence = result.get('whisper_confidence', 0.0)
                transcription.wav2vec_text = result.get('wav2vec_text', '')
                transcription.wav2vec_confidence = result.get('wav2vec_confidence', 0.0)
                transcription.comparison_metrics = result.get('comparison_metrics', {})
                
                # Extract processing times from metadata
                metadata = result.get('metadata', {})
                transcription.whisper_processing_time = metadata.get('whisper_processing_time_seconds', 0)
                transcription.wav2vec_processing_time = metadata.get('wav2vec_processing_time_seconds', 0)
                
                # Calculate which model processed quicker
                transcription.calculate_faster_model()
                
                # Use the better performing result as main text
                if result.get('whisper_confidence', 0) >= result.get('wav2vec_confidence', 0):
                    transcription.text = transcription.whisper_text
                    transcription.confidence_score = transcription.whisper_confidence
                else:
                    transcription.text = transcription.wav2vec_text
                    transcription.confidence_score = transcription.wav2vec_confidence
                
                transcription.word_count = len(transcription.text.split())
                transcription.model_used = 'both'
                transcription.processing_metadata = result.get('metadata', {})
                
                # Store segments from the better performing model
                better_segments = []
                if result.get('whisper_confidence', 0) >= result.get('wav2vec_confidence', 0):
                    better_segments = result.get('whisper_segments', [])
                else:
                    better_segments = result.get('wav2vec_segments', [])
                
                for seg in better_segments:
                    segment = TranscriptionSegment(
                        transcription_id=transcription_id,
                        start_time=seg['start'],
                        end_time=seg['end'],
                        text=seg['text'],
                        confidence=seg.get('confidence'),
                        speaker=seg.get('speaker')
                    )
                    db.session.add(segment)
                        
            else:
                # Process with single model
                on_progress('ai_processing', 35, f'Επεξεργασία με {model}...')
                
                # Run async transcribe_audio method properly
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    if app:
                        with app.app_context():
                            result = loop.run_until_complete(self.ai_client.transcribe_audio(
                                audio_file_path=audio_file.file_path,
                                language=transcription.language,
                                model=model,
                                original_filename=audio_file.original_filename,
                                transcription_id=str(transcription.id)
                            ))
                    else:
                        result = loop.run_until_complete(self.ai_client.transcribe_audio(
                            audio_file_path=audio_file.file_path,
                            language=transcription.language,
                            model=model,
                            original_filename=audio_file.original_filename,
                            transcription_id=str(transcription.id)
                        ))
                finally:
                    loop.close()
                
                # Store single model results
                transcription.text = result['text']
                transcription.confidence_score = result.get('confidence', 0.0)
                transcription.word_count = len(result['text'].split())
                transcription.model_used = model  # Use the user's actual choice, not ASR service implementation
                transcription.processing_metadata = result.get('metadata', {})
                
                # Store model-specific results with processing times
                processing_time_seconds = result.get('metadata', {}).get('processing_time_seconds', 0)
                if model == 'whisper':
                    transcription.whisper_text = result['text']
                    transcription.whisper_confidence = result.get('confidence', 0.0)
                    transcription.whisper_processing_time = processing_time_seconds
                elif model == 'wav2vec2':
                    transcription.wav2vec_text = result['text']
                    transcription.wav2vec_confidence = result.get('confidence', 0.0)
                    transcription.wav2vec_processing_time = processing_time_seconds
                
                # Calculate which model processed quicker (for single model cases)
                transcription.calculate_faster_model()
                
                # Save segments if provided
                if 'segments' in result:
                    for seg in result['segments']:
                        segment = TranscriptionSegment(
                            transcription_id=transcription_id,
                            start_time=seg['start'],
                            end_time=seg['end'],
                            text=seg['text'],
                            confidence=seg.get('confidence'),
                            speaker=seg.get('speaker')
                        )
                        db.session.add(segment)
            
            transcription.credits_used = self._calculate_credits(
                transcription.duration_seconds
            )
            
            current_app.logger.info("Academic mode - recording usage for research analytics only")
            current_app.logger.info(f"Academic transcription completed: {transcription.duration_seconds}s, "
                                  f"{transcription.credits_used} research units processed")
            
            on_progress('finalizing', 95, 'Ολοκλήρωση μεταγραφής...')
            
            transcription.status = 'completed'
            transcription.completed_at = datetime.utcnow()
            
            audio_file.status = 'completed'
            
            db.session.commit()
            
            # Cache the completed transcription in Redis
            try:
                if get_transcription_cache().cache_transcription(transcription):
                    logger.info(f"Cached completed transcription {transcription_id} in Redis")
                else:
                    logger.warning(f"Failed to cache transcription {transcription_id} in Redis")
            except Exception as cache_error:
                logger.warning(f"Redis caching error for transcription {transcription_id}: {str(cache_error)}")
            
            try:
                from app.analytics.services import AnalyticsService
                analytics_service = AnalyticsService()
                analytics_service._invalidate_cache_pattern(f'user_{user_id}')
                analytics_service._invalidate_cache_pattern('user_stats')
                analytics_service._invalidate_cache_pattern('system')
                logger.info(f"Analytics cache invalidated for user {user_id}")
            except Exception as cache_error:
                logger.warning(f"Failed to invalidate analytics cache: {str(cache_error)}")
            
            if progress_manager and socketio:
                try:
                    completion_data = {
                        'transcription_id': str(transcription_id),
                        'text_preview': transcription.text[:200] + '...' if len(transcription.text) > 200 else transcription.text,
                        'word_count': transcription.word_count,
                        'confidence_score': transcription.confidence_score,
                        'duration_seconds': transcription.duration_seconds
                    }
                    if app:
                        with app.app_context():
                            progress_manager.broadcast_completion(str(transcription_id), completion_data)
                    else:
                        progress_manager.broadcast_completion(str(transcription_id), completion_data)
                except Exception as e:
                    logger.error(f"Completion broadcast error | transcription_id={transcription_id} | error={str(e)}")
            
            logger.info(f"Transcription completed successfully | transcription_id={transcription_id}")
            
        except Exception as e:
            transcription.status = 'failed'
            transcription.error_message = str(e)
            transcription.completed_at = datetime.utcnow()
            
            audio_file.status = 'failed'
            
            db.session.commit()
            
            if progress_manager and socketio:
                try:
                    if app:
                        with app.app_context():
                            progress_manager.broadcast_error(
                                str(transcription_id), 
                                f"Σφάλμα κατά την επεξεργασία: {str(e)}", 
                                'TRANSCRIPTION_FAILED'
                            )
                    else:
                        progress_manager.broadcast_error(
                            str(transcription_id), 
                            f"Σφάλμα κατά την επεξεργασία: {str(e)}", 
                            'TRANSCRIPTION_FAILED'
                        )
                except Exception as broadcast_error:
                    logger.error(f"Error broadcast failed | transcription_id={transcription_id} | error={str(broadcast_error)}")
            
            logger.error(f"Transcription failed | transcription_id={transcription_id} | error={str(e)}")
    
    def get_transcription(self, transcription_id: int, user_id: int) -> Optional[Transcription]:
        """Get transcription if user has access. Check Redis cache first for faster retrieval."""
        
        # First check Redis cache for completed transcriptions
        try:
            cached_data = get_transcription_cache().get_cached_transcription(transcription_id, user_id)
            if cached_data:
                # Convert cached data back to Transcription object
                logger.info(f"Retrieved transcription {transcription_id} from Redis cache")
                cached_transcription = self._create_transcription_from_cache(cached_data)
                if cached_transcription:
                    return cached_transcription
                else:
                    logger.warning(f"Failed to reconstruct transcription {transcription_id} from cache, falling back to database")
        except Exception as cache_error:
            logger.warning(f"Redis cache retrieval error for transcription {transcription_id}: {str(cache_error)}")
        
        # Fallback to database
        logger.debug(f"Fetching transcription {transcription_id} from database")
        transcription = self.repository.get_by_id(transcription_id)
        
        if transcription and transcription.user_id == user_id:
            # Cache completed transcriptions for future requests
            if transcription.status == 'completed':
                try:
                    get_transcription_cache().cache_transcription(transcription)
                    logger.debug(f"Cached transcription {transcription_id} after database retrieval")
                except Exception as cache_error:
                    logger.warning(f"Failed to cache transcription after DB fetch: {str(cache_error)}")
            
            return transcription
        
        return None
    
    def _create_transcription_from_cache(self, cached_data: Dict[str, Any]) -> Transcription:
        """Creates a Transcription object from cached data."""
        try:
            transcription = Transcription()
        
            # Basic fields
            transcription.id = cached_data.get('id')
            transcription.title = cached_data.get('title')
            transcription.description = cached_data.get('description')
            transcription.text = cached_data.get('text')
            transcription.language = cached_data.get('language')
            transcription.status = cached_data.get('status')
            transcription.model_used = cached_data.get('model_used')
            transcription.confidence_score = cached_data.get('confidence_score')
            transcription.word_count = cached_data.get('word_count')
            transcription.duration_seconds = cached_data.get('duration_seconds')
            transcription.processing_time = cached_data.get('processing_time')
            transcription.user_id = cached_data.get('user_id')
            transcription.audio_file_id = cached_data.get('audio_file_id')
            transcription.error_message = cached_data.get('error_message')
            transcription.credits_used = cached_data.get('credits_used')
        
            # Comparison data
            transcription.whisper_text = cached_data.get('whisper_text')
            transcription.whisper_confidence = cached_data.get('whisper_confidence')
            transcription.whisper_processing_time = cached_data.get('whisper_processing_time')
            transcription.wav2vec_text = cached_data.get('wav2vec_text')
            transcription.wav2vec_confidence = cached_data.get('wav2vec_confidence')
            transcription.wav2vec_processing_time = cached_data.get('wav2vec_processing_time')
            
            # Performance metrics
            transcription.whisper_wer = cached_data.get('whisper_wer')
            transcription.whisper_cer = cached_data.get('whisper_cer')
            transcription.whisper_accuracy = cached_data.get('whisper_accuracy')
            transcription.wav2vec_wer = cached_data.get('wav2vec_wer')
            transcription.wav2vec_cer = cached_data.get('wav2vec_cer')
            transcription.wav2vec_accuracy = cached_data.get('wav2vec_accuracy')
            transcription.best_performing_model = cached_data.get('best_performing_model')
            transcription.faster_model = cached_data.get('faster_model')
            transcription.academic_accuracy_score = cached_data.get('academic_accuracy_score')
            transcription.evaluation_completed = cached_data.get('evaluation_completed')
            transcription.ground_truth_text = cached_data.get('ground_truth_text')
            
            # Date fields - convert from ISO strings
            if cached_data.get('created_at'):
                transcription.created_at = datetime.fromisoformat(cached_data['created_at'])
            if cached_data.get('completed_at'):
                transcription.completed_at = datetime.fromisoformat(cached_data['completed_at'])
            if cached_data.get('started_at'):
                transcription.started_at = datetime.fromisoformat(cached_data['started_at'])
            if cached_data.get('updated_at'):
                transcription.updated_at = datetime.fromisoformat(cached_data['updated_at'])
            
            # Mark as cached so we know it came from Redis
            transcription._is_from_cache = True
            
            # Load essential relationships from database for API compatibility
            from app.transcription.models import TranscriptionSegment
            from app.audio.models import AudioFile
            
            # Load segments if they exist (but don't attach to session)
            segments = TranscriptionSegment.query.filter_by(transcription_id=transcription.id).all()
            transcription.segments = segments
            
            # Load audio file relationship if it exists
            if transcription.audio_file_id:
                audio_file = AudioFile.query.get(transcription.audio_file_id)
                transcription.audio_file = audio_file
            
            return transcription
            
        except Exception as e:
            logger.error(f"Error creating transcription from cache: {str(e)}")
            # Return None to trigger database fallback
            return None
    
    def get_user_transcriptions(self, user_id: int, page: int = 1, 
                               per_page: int = 20, 
                               status: Optional[str] = None,
                               sort: str = 'created_at',
                               order: str = 'desc',
                               start_date: Optional[str] = None,
                               end_date: Optional[str] = None,
                               search: Optional[str] = None) -> Tuple[List[Transcription], Dict[str, Any]]:
        """Get all transcriptions for a user with pagination and filtering."""
        filters = {'user_id': user_id}
        if status:
            filters['status'] = status
        if start_date:
            filters['start_date'] = start_date  # YYYY-MM-DD format from frontend
        if end_date:
            filters['end_date'] = end_date      # YYYY-MM-DD format from frontend
        if search:
            filters['search'] = search.strip()
        
        result = self.repository.paginate(
            page=page,
            per_page=per_page,
            filters=filters,
            sort_by=sort,
            sort_order=order
        )
        
        return result['items'], {
            'total': result['total'],
            'pages': result['pages'],
            'current_page': result['current_page'],
            'per_page': result['per_page'],
            'has_prev': result['has_prev'],
            'has_next': result['has_next']
        }
    
    def update_transcription_text(self, transcription_id: int, 
                                 user_id: int, text: str) -> Optional[Transcription]:
        """Update transcription text (manual corrections)."""
        transcription = self.get_transcription(transcription_id, user_id)
        
        if not transcription:
            return None
        
        # Invalidate cache before updating
        try:
            get_transcription_cache().invalidate_transcription(transcription_id, user_id)
            logger.debug(f"Invalidated cache for transcription {transcription_id} before text update")
        except Exception as cache_error:
            logger.warning(f"Failed to invalidate cache before update: {str(cache_error)}")
        
        transcription.text = text
        transcription.word_count = len(text.split())
        transcription.updated_at = datetime.utcnow()
        db.session.commit()
        
        # Re-cache the updated transcription if it's completed
        if transcription.status == 'completed':
            try:
                get_transcription_cache().cache_transcription(transcription)
                logger.debug(f"Re-cached updated transcription {transcription_id}")
            except Exception as cache_error:
                logger.warning(f"Failed to re-cache updated transcription: {str(cache_error)}")
        
        return transcription
    
    def delete_transcription(self, transcription_id: int, user_id: int) -> bool:
        """Delete transcription if user has access."""
        transcription = self.get_transcription(transcription_id, user_id)
        
        if not transcription:
            return False
        
        # Invalidate cache before deletion
        try:
            get_transcription_cache().invalidate_transcription(transcription_id, user_id)
            logger.debug(f"Invalidated cache for transcription {transcription_id} before deletion")
        except Exception as cache_error:
            logger.warning(f"Failed to invalidate cache before deletion: {str(cache_error)}")
        
        transcription.soft_delete()
        return True
    
    def retry_transcription(self, transcription_id: int, user_id: int) -> Optional[Transcription]:
        """Retry a failed transcription."""
        transcription = self.get_transcription(transcription_id, user_id)
        
        if not transcription or transcription.status not in ['failed', 'completed']:
            return None
        
        # Reset status
        transcription.status = 'pending'
        transcription.error_message = None
        transcription.started_at = None
        transcription.completed_at = None
        db.session.commit()
        
        # Start processing again
        self._process_transcription_async(transcription.id)
        
        return transcription
    
    def export_transcription(self, transcription_id: int, user_id: int, 
                           format_type: str) -> Tuple[bytes, str, str]:
        """Export transcription in various formats."""
        transcription = self.get_transcription(transcription_id, user_id)
        
        if not transcription:
            return None, None, None
        
        if format_type == 'txt':
            return self._export_as_txt(transcription)
        elif format_type == 'srt':
            return self._export_as_srt(transcription)
        elif format_type == 'docx':
            return self._export_as_docx(transcription)
        elif format_type == 'pdf':
            return self._export_as_pdf(transcription)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def _export_as_txt(self, transcription: Transcription) -> Tuple[bytes, str, str]:
        """Export as plain text."""
        
        # Check if this is comparison mode with both models
        is_comparison = (transcription.model_used == 'both' and 
                        transcription.whisper_text and 
                        transcription.wav2vec_text)
        
        if is_comparison:
            # Build comprehensive comparison content
            content_lines = []
            content_lines.append(f"ΣΥΓΚΡΙΤΙΚΗ ΜΕΤΑΓΡΑΦΗ")
            content_lines.append(f"Τίτλος: {transcription.title}")
            content_lines.append(f"Ημερομηνία: {transcription.created_at.strftime('%d/%m/%Y %H:%M')}")
            content_lines.append(f"Διάρκεια: {transcription.duration_seconds:.2f} δευτερόλεπτα")
            content_lines.append("")
            content_lines.append("=" * 80)
            content_lines.append("")
            
            # Whisper results
            content_lines.append("ΑΠΟΤΕΛΕΣΜΑ WHISPER")
            content_lines.append("-" * 50)
            content_lines.append(transcription.whisper_text)
            content_lines.append("")
            
            if transcription.whisper_confidence:
                content_lines.append(f"Εμπιστοσύνη: {transcription.whisper_confidence:.2f}")
            if transcription.whisper_processing_time:
                content_lines.append(f"Χρόνος επεξεργασίας: {transcription.whisper_processing_time:.2f}s")
            if transcription.whisper_wer is not None:
                content_lines.append(f"WER: {transcription.whisper_wer:.2f}%")
            if transcription.whisper_cer is not None:
                content_lines.append(f"CER: {transcription.whisper_cer:.2f}%")
            if transcription.whisper_accuracy is not None:
                content_lines.append(f"Ακρίβεια: {transcription.whisper_accuracy:.2f}%")
            
            content_lines.append("")
            content_lines.append("=" * 80)
            content_lines.append("")
            
            # Wav2Vec2 results
            content_lines.append("ΑΠΟΤΕΛΕΣΜΑ WAV2VEC2")
            content_lines.append("-" * 50)
            content_lines.append(transcription.wav2vec_text)
            content_lines.append("")
            
            if transcription.wav2vec_confidence:
                content_lines.append(f"Εμπιστοσύνη: {transcription.wav2vec_confidence:.2f}")
            if transcription.wav2vec_processing_time:
                content_lines.append(f"Χρόνος επεξεργασίας: {transcription.wav2vec_processing_time:.2f}s")
            if transcription.wav2vec_wer is not None:
                content_lines.append(f"WER: {transcription.wav2vec_wer:.2f}%")
            if transcription.wav2vec_cer is not None:
                content_lines.append(f"CER: {transcription.wav2vec_cer:.2f}%")
            if transcription.wav2vec_accuracy is not None:
                content_lines.append(f"Ακρίβεια: {transcription.wav2vec_accuracy:.2f}%")
            
            # Comparison summary if evaluation exists
            if (transcription.evaluation_completed and 
                transcription.whisper_wer is not None and 
                transcription.wav2vec_wer is not None):
                content_lines.append("")
                content_lines.append("=" * 80)
                content_lines.append("")
                content_lines.append("ΣΥΓΚΡΙΤΙΚΗ ΑΝΑΛΥΣΗ")
                content_lines.append("-" * 50)
                content_lines.append(f"Καλύτερο μοντέλο (ακρίβεια): {transcription.best_performing_model or 'N/A'}")
                content_lines.append(f"Ταχύτερο μοντέλο: {transcription.faster_model or 'N/A'}")
                content_lines.append(f"Βαθμολογία ακρίβειας: {transcription.academic_accuracy_score or 0:.2f}%")
                
                if transcription.ground_truth_text:
                    content_lines.append("")
                    content_lines.append("ΣΩΣΤΟ ΚΕΙΜΕΝΟ (Ground Truth)")
                    content_lines.append("-" * 50)
                    content_lines.append(transcription.ground_truth_text)
            
            content = '\n'.join(content_lines)
            filename = f"comparison_transcription_{transcription.id}.txt"
        else:
            # Single model export
            content = transcription.text
            filename = f"transcription_{transcription.id}.txt"
        
        return content.encode('utf-8'), filename, 'text/plain'
    
    def _export_as_srt(self, transcription: Transcription) -> Tuple[bytes, str, str]:
        """Export as SRT subtitle file."""
        srt_content = []
        
        if transcription.segments and len(transcription.segments) > 0:
            # Use actual segments if available
            for i, segment in enumerate(transcription.segments, 1):
                start_time = self._seconds_to_srt_time(segment.start_time)
                end_time = self._seconds_to_srt_time(segment.end_time)
                
                srt_content.append(f"{i}")
                srt_content.append(f"{start_time} --> {end_time}")
                srt_content.append(segment.text)
                srt_content.append("")  # Empty line between subtitles
        else:
            # Fallback: Create basic segments from full text
            logger.warning(f"No segments found for transcription {transcription.id}, creating fallback SRT")
            
            # Split text into sentences for basic segmentation
            text = transcription.text if transcription.text else "Δεν υπάρχει κείμενο μεταγραφής"
            sentences = self._split_text_for_srt(text)
            
            duration = transcription.duration_seconds if transcription.duration_seconds else 30.0
            segment_duration = duration / len(sentences) if sentences else duration
            
            for i, sentence in enumerate(sentences, 1):
                start_time = (i - 1) * segment_duration
                end_time = i * segment_duration
                
                start_srt = self._seconds_to_srt_time(start_time)
                end_srt = self._seconds_to_srt_time(end_time)
                
                srt_content.append(f"{i}")
                srt_content.append(f"{start_srt} --> {end_srt}")
                srt_content.append(sentence.strip())
                srt_content.append("")  # Empty line between subtitles
        
        content = '\n'.join(srt_content)
        filename = f"transcription_{transcription.id}.srt"
        
        return content.encode('utf-8'), filename, 'application/x-subrip'
    
    def _export_as_docx(self, transcription: Transcription) -> Tuple[bytes, str, str]:
        """Export as DOCX document."""
        try:
            from docx import Document
            from docx.shared import Inches
            from io import BytesIO
            
            # Check if this is comparison mode
            is_comparison = (transcription.model_used == 'both' and 
                            transcription.whisper_text and 
                            transcription.wav2vec_text)
            
            # Create document
            doc = Document()
            
            if is_comparison:
                # Add title for comparison
                title = doc.add_heading(f'Συγκριτική Μεταγραφή: {transcription.title}', 0)
                
                # Add metadata
                doc.add_paragraph(f'Ημερομηνία: {transcription.created_at.strftime("%d/%m/%Y %H:%M")}')
                doc.add_paragraph(f'Διάρκεια: {transcription.duration_seconds:.2f} δευτερόλεπτα')
                doc.add_paragraph(f'Μοντέλο: Σύγκριση Whisper vs Wav2Vec2')
                
                # Add separator
                doc.add_paragraph('═' * 60)
                
                # Whisper results section
                doc.add_heading('Αποτέλεσμα Whisper', level=1)
                doc.add_paragraph(transcription.whisper_text)
                
                # Whisper metrics
                doc.add_heading('Μετρικές Whisper', level=2)
                if transcription.whisper_confidence:
                    doc.add_paragraph(f'Εμπιστοσύνη: {transcription.whisper_confidence:.2f}')
                if transcription.whisper_processing_time:
                    doc.add_paragraph(f'Χρόνος επεξεργασίας: {transcription.whisper_processing_time:.2f}s')
                if transcription.whisper_wer is not None:
                    doc.add_paragraph(f'WER: {transcription.whisper_wer:.2f}%')
                if transcription.whisper_cer is not None:
                    doc.add_paragraph(f'CER: {transcription.whisper_cer:.2f}%')
                if transcription.whisper_accuracy is not None:
                    doc.add_paragraph(f'Ακρίβεια: {transcription.whisper_accuracy:.2f}%')
                
                doc.add_paragraph('═' * 60)
                
                # Wav2Vec2 results section
                doc.add_heading('Αποτέλεσμα Wav2Vec2', level=1)
                doc.add_paragraph(transcription.wav2vec_text)
                
                # Wav2Vec2 metrics
                doc.add_heading('Μετρικές Wav2Vec2', level=2)
                if transcription.wav2vec_confidence:
                    doc.add_paragraph(f'Εμπιστοσύνη: {transcription.wav2vec_confidence:.2f}')
                if transcription.wav2vec_processing_time:
                    doc.add_paragraph(f'Χρόνος επεξεργασίας: {transcription.wav2vec_processing_time:.2f}s')
                if transcription.wav2vec_wer is not None:
                    doc.add_paragraph(f'WER: {transcription.wav2vec_wer:.2f}%')
                if transcription.wav2vec_cer is not None:
                    doc.add_paragraph(f'CER: {transcription.wav2vec_cer:.2f}%')
                if transcription.wav2vec_accuracy is not None:
                    doc.add_paragraph(f'Ακρίβεια: {transcription.wav2vec_accuracy:.2f}%')
                
                # Comparison analysis if available
                if (transcription.evaluation_completed and 
                    transcription.whisper_wer is not None and 
                    transcription.wav2vec_wer is not None):
                    
                    doc.add_paragraph('═' * 60)
                    doc.add_heading('Συγκριτική Ανάλυση', level=1)
                    doc.add_paragraph(f'Καλύτερο μοντέλο (ακρίβεια): {transcription.best_performing_model or "N/A"}')
                    doc.add_paragraph(f'Ταχύτερο μοντέλο: {transcription.faster_model or "N/A"}')
                    doc.add_paragraph(f'Βαθμολογία ακρίβειας: {transcription.academic_accuracy_score or 0:.2f}%')
                    
                    if transcription.ground_truth_text:
                        doc.add_heading('Σωστό Κείμενο (Ground Truth)', level=2)
                        doc.add_paragraph(transcription.ground_truth_text)
                
                filename = f"comparison_transcription_{transcription.id}.docx"
            else:
                # Single model document
                title = doc.add_heading(f'Μεταγραφή: {transcription.title}', 0)
                
                # Add metadata
                doc.add_paragraph(f'Ημερομηνία: {transcription.created_at.strftime("%d/%m/%Y %H:%M")}')
                doc.add_paragraph(f'Διάρκεια: {transcription.duration_seconds} δευτερόλεπτα')
                if transcription.confidence_score:
                    doc.add_paragraph(f'Εμπιστοσύνη: {transcription.confidence_score}')
                doc.add_paragraph(f'Μοντέλο: {transcription.model_used or "N/A"}')
                
                # Add separator
                doc.add_paragraph('─' * 50)
                
                # Add main text
                doc.add_heading('Κείμενο Μεταγραφής', level=2)
                doc.add_paragraph(transcription.text)
                
                filename = f"transcription_{transcription.id}.docx"
            
            # Add segments if available
            if transcription.segments:
                doc.add_heading('Τμήματα με Χρονικές Σφραγίδες', level=2)
                for i, segment in enumerate(transcription.segments, 1):
                    start_time = f"{int(segment.start_time // 60):02d}:{int(segment.start_time % 60):02d}"
                    end_time = f"{int(segment.end_time // 60):02d}:{int(segment.end_time % 60):02d}"
                    doc.add_paragraph(f'[{start_time} - {end_time}] {segment.text}')
            
            # Save to BytesIO
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            return docx_buffer.read(), filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
        except ImportError:
            logger.warning("python-docx not installed, falling back to text export")
            # Fallback to text export
            content = f"Μεταγραφή: {transcription.title}\n"
            content += f"Ημερομηνία: {transcription.created_at.strftime('%d/%m/%Y %H:%M')}\n"
            content += f"Διάρκεια: {transcription.duration_seconds} δευτερόλεπτα\n"
            if transcription.confidence_score:
                content += f"Εμπιστοσύνη: {transcription.confidence_score}\n"
            content += f"Μοντέλο: {transcription.model_used or 'N/A'}\n"
            content += "\n" + "─" * 50 + "\n\n"
            content += transcription.text
            
            filename = f"transcription_{transcription.id}.docx"
            return content.encode('utf-8'), filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        except Exception as e:
            logger.error(f"Error creating DOCX: {str(e)}")
            # Fallback to text export
            content = f"Μεταγραφή: {transcription.title}\n"
            content += f"Ημερομηνία: {transcription.created_at.strftime('%d/%m/%Y %H:%M')}\n"
            content += f"Διάρκεια: {transcription.duration_seconds} δευτερόλεπτα\n"
            if transcription.confidence_score:
                content += f"Εμπιστοσύνη: {transcription.confidence_score}\n"
            content += f"Μοντέλο: {transcription.model_used or 'N/A'}\n"
            content += "\n" + "─" * 50 + "\n\n"
            content += transcription.text
            
            filename = f"transcription_{transcription.id}.docx"
            return content.encode('utf-8'), filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    def _export_as_pdf(self, transcription: Transcription) -> Tuple[bytes, str, str]:
        """Export as PDF document."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from io import BytesIO
            import os
            
            # Create PDF buffer
            pdf_buffer = BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Try to register a font that supports Greek characters
            try:
                # Use DejaVu Sans if available (supports Greek)
                font_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
                    greek_font = 'DejaVuSans'
                else:
                    greek_font = 'Helvetica'  # Fallback
            except:
                greek_font = 'Helvetica'  # Fallback
            
            # Custom styles for Greek text
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontName=greek_font,
                fontSize=18,
                textColor='black',
                alignment=TA_CENTER,
                spaceAfter=30
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontName=greek_font,
                fontSize=14,
                textColor='black',
                spaceBefore=20,
                spaceAfter=10
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontName=greek_font,
                fontSize=10,
                textColor='black',
                leftIndent=0,
                rightIndent=0,
                spaceAfter=10,
                alignment=TA_LEFT
            )
            
            # Build PDF content
            content = []
            
            # Title
            content.append(Paragraph(f'Μεταγραφή: {transcription.title}', title_style))
            content.append(Spacer(1, 20))
            
            # Metadata
            metadata = f"""<b>Ημερομηνία:</b> {transcription.created_at.strftime('%d/%m/%Y %H:%M')}<br/>
            <b>Διάρκεια:</b> {transcription.duration_seconds} δευτερόλεπτα<br/>
            <b>Μοντέλο:</b> {transcription.model_used or 'N/A'}<br/>"""
            
            if transcription.confidence_score:
                metadata += f"<b>Εμπιστοσύνη:</b> {transcription.confidence_score}<br/>"
            
            content.append(Paragraph(metadata, body_style))
            content.append(Spacer(1, 20))
            
            # Main text
            content.append(Paragraph('Κείμενο Μεταγραφής', heading_style))
            
            # Split text into paragraphs for better PDF formatting
            text_paragraphs = transcription.text.split('\n')
            for para in text_paragraphs:
                if para.strip():
                    content.append(Paragraph(para.strip(), body_style))
            
            # Add segments if available
            if transcription.segments:
                content.append(Spacer(1, 20))
                content.append(Paragraph('Τμήματα με Χρονικές Σφραγίδες', heading_style))
                
                for i, segment in enumerate(transcription.segments, 1):
                    start_time = f"{int(segment.start_time // 60):02d}:{int(segment.start_time % 60):02d}"
                    end_time = f"{int(segment.end_time // 60):02d}:{int(segment.end_time % 60):02d}"
                    segment_text = f"[{start_time} - {end_time}] {segment.text}"
                    content.append(Paragraph(segment_text, body_style))
            
            # Build PDF
            doc.build(content)
            pdf_buffer.seek(0)
            
            filename = f"transcription_{transcription.id}.pdf"
            return pdf_buffer.read(), filename, 'application/pdf'
            
        except ImportError:
            logger.warning("reportlab not installed, falling back to text export")
            # Fallback to text export
            content = f"Μεταγραφή: {transcription.title}\n"
            content += f"Ημερομηνία: {transcription.created_at.strftime('%d/%m/%Y %H:%M')}\n"
            content += f"Διάρκεια: {transcription.duration_seconds} δευτερόλεπτα\n"
            if transcription.confidence_score:
                content += f"Εμπιστοσύνη: {transcription.confidence_score}\n"
            content += f"Μοντέλο: {transcription.model_used or 'N/A'}\n"
            content += "\n" + "─" * 50 + "\n\n"
            content += transcription.text
            
            filename = f"transcription_{transcription.id}.pdf"
            return content.encode('utf-8'), filename, 'application/pdf'
        except Exception as e:
            logger.error(f"Error creating PDF: {str(e)}")
            # Fallback to text export
            content = f"Μεταγραφή: {transcription.title}\n"
            content += f"Ημερομηνία: {transcription.created_at.strftime('%d/%m/%Y %H:%M')}\n"
            content += f"Διάρκεια: {transcription.duration_seconds} δευτερόλεπτα\n"
            if transcription.confidence_score:
                content += f"Εμπιστοσύνη: {transcription.confidence_score}\n"
            content += f"Μοντέλο: {transcription.model_used or 'N/A'}\n"
            content += "\n" + "─" * 50 + "\n\n"
            content += transcription.text
            
            filename = f"transcription_{transcription.id}.pdf"
            return content.encode('utf-8'), filename, 'application/pdf'
    
    def _seconds_to_srt_time(self, seconds: float) -> str:
        """Convert seconds to SRT time format (HH:MM:SS,mmm)."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
    
    def _split_text_for_srt(self, text: str) -> list:
        """Split text into manageable segments for SRT subtitles."""
        if not text or len(text.strip()) == 0:
            return ["Δεν υπάρχει κείμενο"]
        
        # Try to split by sentences first (Greek punctuation aware)
        import re
        sentences = re.split(r'[.!;·]\s+', text.strip())
        
        # If sentences are too long (>80 chars), split further
        result = []
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            if len(sentence) <= 80:
                result.append(sentence)
            else:
                # Split long sentences by clauses or phrases
                clauses = re.split(r'[,·]\s+', sentence)
                current_segment = ""
                
                for clause in clauses:
                    clause = clause.strip()
                    if len(current_segment + " " + clause) <= 80:
                        current_segment = current_segment + " " + clause if current_segment else clause
                    else:
                        if current_segment:
                            result.append(current_segment)
                        current_segment = clause
                
                if current_segment:
                    result.append(current_segment)
        
        # Ensure we have at least one segment
        if not result:
            result = [text[:80] + "..." if len(text) > 80 else text]
        
        return result
    
    def _calculate_credits(self, duration_seconds: float) -> float:
        """Calculate credits used based on duration."""
        return duration_seconds / 60.0
    
    def generate_ai_summary(self, transcription_id: int, user_id: int) -> Optional[str]:
        """Generate simple summary for transcription."""
        transcription = self.get_transcription(transcription_id, user_id)
        
        if not transcription or transcription.status != 'completed':
            return None
        
        try:
            # Use simple summarization approach
            logger.info(f"Generating simple summary for transcription {transcription_id}")
            text = transcription.text
            
            # Basic Greek text summarization
            sentences = text.split('. ')
            
            if len(sentences) <= 3:
                return text  # Already short enough
            
            # Take first and last sentences plus any containing key terms
            key_terms = ['σημαντικό', 'κύριο', 'βασικό', 'ουσιαστικό', 'προβλήματα', 'λύσεις']
            
            summary_sentences = [sentences[0]]  # First sentence
            
            # Add sentences with key terms
            for sentence in sentences[1:-1]:
                if any(term in sentence.lower() for term in key_terms):
                    summary_sentences.append(sentence)
                    if len(summary_sentences) >= 5:  # Limit summary length
                        break
            
            # Add last sentence if not already included
            if len(sentences) > 1 and sentences[-1] not in summary_sentences:
                summary_sentences.append(sentences[-1])
            
            summary = '. '.join(summary_sentences)
            
            # Add summary metadata
            summary_html = f"""<div class="ai-summary">
                <h4>Κύρια Σημεία:</h4>
                <p>{summary}</p>
                <hr>
                <small class="summary-stats">
                    Αρχικό κείμενο: {len(text.split())} λέξεις | 
                    Περίληψη: {len(summary.split())} λέξεις | 
                    Συμπίεση: {int((1 - len(summary.split()) / len(text.split())) * 100)}%
                </small>
            </div>"""
            
            return summary_html
            
        except Exception as e:
            logger.error(f"Error generating summary: {str(e)}")
            return None
    
    def update_segment(self, transcription_id: int, segment_index: int, text: str, user_id: int) -> bool:
        """Update a specific transcription segment."""
        transcription = self.get_transcription(transcription_id, user_id)
        
        if not transcription:
            return False
        
        try:
            if segment_index < 0 or segment_index >= len(transcription.segments):
                return False
            
            segment = transcription.segments[segment_index]
            segment.text = text
            
            # Update the main transcription text
            updated_text = ' '.join([seg.text for seg in transcription.segments])
            transcription.text = updated_text
            transcription.word_count = len(updated_text.split())
            
            db.session.commit()
            return True
            
        except Exception as e:
            logger.error(f"Error updating segment: {str(e)}")
            db.session.rollback()
            return False

    # =============================================================================
    # BULK OPERATIONS FOR WEEK 3
    # =============================================================================

    def bulk_update_segments(self, user_id: int, transcription_id: int, updates: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Bulk update transcription segments."""
        try:
            # Get transcription and verify ownership
            transcription = Transcription.query.filter_by(id=transcription_id, user_id=user_id).first()
            if not transcription:
                raise ValueError("Transcription not found or access denied")
            
            updated_count = 0
            errors = []
            
            for update in updates:
                try:
                    segment_id = update.get('id')
                    segment = TranscriptionSegment.query.filter_by(
                        id=segment_id, 
                        transcription_id=transcription_id
                    ).first()
                    
                    if not segment:
                        errors.append(f"Segment {segment_id} not found")
                        continue
                    
                    # Update segment fields
                    if 'text' in update:
                        segment.text = update['text']
                    if 'confidence' in update:
                        segment.confidence = update['confidence']
                    if 'speaker' in update:
                        segment.speaker = update['speaker']
                    if 'start_time' in update:
                        segment.start_time = update['start_time']
                    if 'end_time' in update:
                        segment.end_time = update['end_time']
                    
                    updated_count += 1
                    
                except Exception as e:
                    errors.append(f"Error updating segment {update.get('id', 'unknown')}: {str(e)}")
            
            # Update main transcription text if segments were updated
            if updated_count > 0:
                updated_text = ' '.join([seg.text for seg in transcription.segments])
                transcription.text = updated_text
                transcription.word_count = len(updated_text.split())
                transcription.updated_at = datetime.utcnow()
            
            db.session.commit()
            
            return {
                'updated_count': updated_count,
                'total_requested': len(updates),
                'errors': errors,
                'transcription_id': transcription_id
            }
            
        except Exception as e:
            db.session.rollback()
            logger.error(f"Bulk segment update failed: {str(e)}")
            raise

    def bulk_delete_transcriptions(self, user_id: int, transcription_ids: List[int], permanent: bool = False) -> Dict[str, Any]:
        """Bulk delete transcriptions."""
        try:
            # Get transcriptions owned by user
            transcriptions = Transcription.query.filter(
                Transcription.id.in_(transcription_ids),
                Transcription.user_id == user_id
            ).all()
            
            found_ids = [t.id for t in transcriptions]
            not_found_ids = [tid for tid in transcription_ids if tid not in found_ids]
            
            deleted_count = 0
            
            for transcription in transcriptions:
                if permanent:
                    # Hard delete
                    db.session.delete(transcription)
                else:
                    # Soft delete
                    transcription.is_deleted = True
                    transcription.deleted_at = datetime.utcnow()
                
                deleted_count += 1
            
            db.session.commit()
            
            return {
                'deleted_count': deleted_count,
                'not_found_count': len(not_found_ids),
                'not_found_ids': not_found_ids,
                'permanent': permanent
            }
            
        except Exception as e:
            db.session.rollback()
            logger.error(f"Bulk delete failed: {str(e)}")
            raise

    def bulk_update_status(self, user_id: int, transcription_ids: List[int], status: str, reason: str) -> Dict[str, Any]:
        """Bulk update transcription status."""
        try:
            # Get transcriptions owned by user
            transcriptions = Transcription.query.filter(
                Transcription.id.in_(transcription_ids),
                Transcription.user_id == user_id
            ).all()
            
            found_ids = [t.id for t in transcriptions]
            not_found_ids = [tid for tid in transcription_ids if tid not in found_ids]
            
            updated_count = 0
            
            for transcription in transcriptions:
                old_status = transcription.status
                transcription.status = status
                transcription.updated_at = datetime.utcnow()
                
                # Log status change reason
                if hasattr(transcription, 'status_history'):
                    transcription.status_history = transcription.status_history or []
                    transcription.status_history.append({
                        'from': old_status,
                        'to': status,
                        'reason': reason,
                        'timestamp': datetime.utcnow().isoformat()
                    })
                
                updated_count += 1
            
            db.session.commit()
            
            return {
                'updated_count': updated_count,
                'not_found_count': len(not_found_ids),
                'not_found_ids': not_found_ids,
                'new_status': status
            }
            
        except Exception as e:
            db.session.rollback()
            logger.error(f"Bulk status update failed: {str(e)}")
            raise

    def bulk_export_transcriptions(self, user_id: int, transcription_ids: List[int], 
                                  format: str, include_segments: bool, include_metadata: bool) -> Dict[str, Any]:
        """Bulk export transcriptions in various formats."""
        try:
            # Get transcriptions owned by user
            transcriptions = Transcription.query.filter(
                Transcription.id.in_(transcription_ids),
                Transcription.user_id == user_id
            ).all()
            
            exported_data = []
            
            for transcription in transcriptions:
                data = {
                    'id': transcription.id,
                    'title': transcription.title,
                    'text': transcription.text,
                    'language': transcription.language,
                    'model_used': transcription.model_used,
                    'status': transcription.status,
                    'created_at': transcription.created_at.isoformat()
                }
                
                if include_metadata:
                    data.update({
                        'description': transcription.description,
                        'duration_seconds': transcription.duration_seconds,
                        'word_count': transcription.word_count,
                        'confidence_score': transcription.confidence_score,
                        'processing_time': transcription.processing_time,
                        'accuracy_score': transcription.accuracy_score
                    })
                
                if include_segments and transcription.segments:
                    data['segments'] = [seg.to_dict() for seg in transcription.segments]
                
                exported_data.append(data)
            
            # Format the export based on requested format
            if format == 'json':
                export_content = exported_data
            elif format == 'csv':
                export_content = self._format_as_csv(exported_data, include_segments)
            elif format == 'txt':
                export_content = self._format_as_txt(exported_data)
            elif format == 'srt':
                export_content = self._format_as_srt(exported_data)
            elif format == 'vtt':
                export_content = self._format_as_vtt(exported_data)
            else:
                export_content = exported_data
            
            return {
                'export_data': export_content,
                'format': format,
                'transcription_count': len(transcriptions),
                'includes_segments': include_segments,
                'includes_metadata': include_metadata
            }
            
        except Exception as e:
            logger.error(f"Bulk export failed: {str(e)}")
            raise

    def bulk_reprocess_transcriptions(self, user_id: int, transcription_ids: List[int], 
                                    model: str, language: str = None, settings: Dict[str, Any] = None) -> Dict[str, Any]:
        """Bulk reprocess transcriptions with different models or settings."""
        try:
            # Get transcriptions owned by user
            transcriptions = Transcription.query.filter(
                Transcription.id.in_(transcription_ids),
                Transcription.user_id == user_id
            ).all()
            
            found_ids = [t.id for t in transcriptions]
            not_found_ids = [tid for tid in transcription_ids if tid not in found_ids]
            
            queued_count = 0
            errors = []
            
            for transcription in transcriptions:
                try:
                    # Update transcription for reprocessing
                    transcription.model_used = model
                    if language:
                        transcription.language = language
                    transcription.status = 'pending'
                    transcription.text = None  # Clear old results
                    transcription.confidence_score = None
                    transcription.processing_time = None
                    transcription.updated_at = datetime.utcnow()
                    
                    # Clear old segments
                    TranscriptionSegment.query.filter_by(transcription_id=transcription.id).delete()
                    
                    # In a real implementation, would queue for processing
                    # For now, just mark as pending
                    queued_count += 1
                    
                except Exception as e:
                    errors.append(f"Error queuing transcription {transcription.id}: {str(e)}")
            
            db.session.commit()
            
            return {
                'queued_count': queued_count,
                'not_found_count': len(not_found_ids),
                'not_found_ids': not_found_ids,
                'errors': errors,
                'model': model,
                'language': language
            }
            
        except Exception as e:
            db.session.rollback()
            logger.error(f"Bulk reprocess failed: {str(e)}")
            raise

    def bulk_analyze_transcriptions(self, user_id: int, transcription_ids: List[int], 
                                  analysis_type: str, group_by: str) -> Dict[str, Any]:
        """Bulk analysis of transcriptions for research purposes."""
        try:
            # Get transcriptions owned by user
            transcriptions = Transcription.query.filter(
                Transcription.id.in_(transcription_ids),
                Transcription.user_id == user_id,
                Transcription.status == 'completed'
            ).all()
            
            analysis_results = {}
            
            if analysis_type == 'accuracy':
                analysis_results = self._analyze_accuracy(transcriptions, group_by)
            elif analysis_type == 'speed':
                analysis_results = self._analyze_speed(transcriptions, group_by)
            elif analysis_type == 'linguistic':
                analysis_results = self._analyze_linguistic(transcriptions, group_by)
            elif analysis_type == 'comparative':
                analysis_results = self._analyze_comparative(transcriptions, group_by)
            
            return {
                'analysis_type': analysis_type,
                'group_by': group_by,
                'analyzed_count': len(transcriptions),
                'results': analysis_results
            }
            
        except Exception as e:
            logger.error(f"Bulk analysis failed: {str(e)}")
            raise

    # Helper methods for export formatting
    
    def _format_as_csv(self, data: List[Dict], include_segments: bool) -> str:
        """Format data as CSV."""
        import csv
        import io
        
        output = io.StringIO()
        if not data:
            return ""
        
        # Basic CSV with main transcription data
        fieldnames = ['id', 'title', 'text', 'language', 'model_used', 'status', 'created_at']
        writer = csv.DictWriter(output, fieldnames=fieldnames)
        writer.writeheader()
        
        for item in data:
            row = {k: v for k, v in item.items() if k in fieldnames}
            writer.writerow(row)
        
        return output.getvalue()
    
    def _format_as_txt(self, data: List[Dict]) -> str:
        """Format data as plain text."""
        output = []
        for item in data:
            output.append(f"Title: {item.get('title', 'Untitled')}")
            output.append(f"Language: {item.get('language', 'Unknown')}")
            output.append(f"Model: {item.get('model_used', 'Unknown')}")
            output.append(f"Text: {item.get('text', '')}")
            output.append("-" * 50)
        
        return "\n".join(output)
    
    def _format_as_srt(self, data: List[Dict]) -> str:
        """Format data as SRT subtitle format."""
        output = []
        subtitle_number = 1
        
        for item in data:
            if 'segments' in item:
                for segment in item['segments']:
                    start_time = self._seconds_to_srt_time(segment['start_time'])
                    end_time = self._seconds_to_srt_time(segment['end_time'])
                    
                    output.append(str(subtitle_number))
                    output.append(f"{start_time} --> {end_time}")
                    output.append(segment['text'])
                    output.append("")
                    subtitle_number += 1
            else:
                # Fallback for transcriptions without segments
                output.append(str(subtitle_number))
                output.append("00:00:00,000 --> 00:00:30,000")
                output.append(item.get('text', ''))
                output.append("")
                subtitle_number += 1
        
        return "\n".join(output)
    
    def _format_as_vtt(self, data: List[Dict]) -> str:
        """Format data as WebVTT format."""
        output = ["WEBVTT", ""]
        
        for item in data:
            if 'segments' in item:
                for segment in item['segments']:
                    start_time = self._seconds_to_vtt_time(segment['start_time'])
                    end_time = self._seconds_to_vtt_time(segment['end_time'])
                    
                    output.append(f"{start_time} --> {end_time}")
                    output.append(segment['text'])
                    output.append("")
            else:
                # Fallback for transcriptions without segments
                output.append("00:00:00.000 --> 00:00:30.000")
                output.append(item.get('text', ''))
                output.append("")
        
        return "\n".join(output)
    
    def _seconds_to_srt_time(self, seconds: float) -> str:
        """Convert seconds to SRT time format."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = seconds % 60
        return f"{hours:02d}:{minutes:02d}:{secs:06.3f}".replace('.', ',')
    
    def _seconds_to_vtt_time(self, seconds: float) -> str:
        """Convert seconds to WebVTT time format."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = seconds % 60
        return f"{hours:02d}:{minutes:02d}:{secs:06.3f}"

    # Helper methods for analysis
    
    def _analyze_accuracy(self, transcriptions: List[Transcription], group_by: str) -> Dict[str, Any]:
        """Analyze accuracy metrics."""
        groups = {}
        
        for trans in transcriptions:
            if trans.accuracy_score is None:
                continue
                
            key = getattr(trans, group_by) if hasattr(trans, group_by) else 'unknown'
            if key not in groups:
                groups[key] = []
            groups[key].append(trans.accuracy_score)
        
        results = {}
        for group, scores in groups.items():
            if scores:
                results[group] = {
                    'mean': sum(scores) / len(scores),
                    'min': min(scores),
                    'max': max(scores),
                    'count': len(scores)
                }
        
        return results
    
    def _analyze_speed(self, transcriptions: List[Transcription], group_by: str) -> Dict[str, Any]:
        """Analyze processing speed metrics."""
        groups = {}
        
        for trans in transcriptions:
            if trans.processing_time is None:
                continue
                
            key = getattr(trans, group_by) if hasattr(trans, group_by) else 'unknown'
            if key not in groups:
                groups[key] = []
            groups[key].append(trans.processing_time)
        
        results = {}
        for group, times in groups.items():
            if times:
                results[group] = {
                    'mean': sum(times) / len(times),
                    'min': min(times),
                    'max': max(times),
                    'count': len(times)
                }
        
        return results
    
    def _analyze_linguistic(self, transcriptions: List[Transcription], group_by: str) -> Dict[str, Any]:
        """Analyze linguistic features."""
        # Simplified linguistic analysis
        groups = {}
        
        for trans in transcriptions:
            key = getattr(trans, group_by) if hasattr(trans, group_by) else 'unknown'
            if key not in groups:
                groups[key] = {
                    'word_count_total': 0,
                    'transcription_count': 0,
                    'languages': set()
                }
            
            groups[key]['word_count_total'] += trans.word_count or 0
            groups[key]['transcription_count'] += 1
            groups[key]['languages'].add(trans.language)
        
        # Convert sets to lists for JSON serialization
        for group in groups.values():
            group['languages'] = list(group['languages'])
            group['avg_words_per_transcription'] = group['word_count_total'] / group['transcription_count'] if group['transcription_count'] > 0 else 0
        
        return groups
    
    def _analyze_comparative(self, transcriptions: List[Transcription], group_by: str) -> Dict[str, Any]:
        """Perform comparative analysis between models."""
        model_groups = {}
        
        for trans in transcriptions:
            model = trans.model_used
            if model not in model_groups:
                model_groups[model] = []
            model_groups[model].append(trans)
        
        # Compare models
        comparison = {}
        for model, trans_list in model_groups.items():
            accuracy_scores = [t.accuracy_score for t in trans_list if t.accuracy_score]
            processing_times = [t.processing_time for t in trans_list if t.processing_time]
            
            comparison[model] = {
                'count': len(trans_list),
                'avg_accuracy': sum(accuracy_scores) / len(accuracy_scores) if accuracy_scores else None,
                'avg_processing_time': sum(processing_times) / len(processing_times) if processing_times else None
            }
        
        return comparison