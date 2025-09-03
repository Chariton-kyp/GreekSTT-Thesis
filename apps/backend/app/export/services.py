import io
from datetime import datetime
from typing import Dict, List, Any, Optional
import pytz
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.transcription.models import Transcription
from app.comparison.models import ModelComparison
from app.utils.correlation_logger import get_correlation_logger

logger = get_correlation_logger(__name__)

def get_greek_time() -> datetime:
    """Get current time in Greek timezone (Europe/Athens)."""
    greek_tz = pytz.timezone('Europe/Athens')
    return datetime.now(greek_tz)

def format_greek_datetime(dt: datetime = None, format_str: str = '%d/%m/%Y %H:%M') -> str:
    """Format datetime in Greek timezone with specified format."""
    if dt is None:
        dt = get_greek_time()
    elif dt.tzinfo is None:
        # If datetime is naive, assume it's UTC and convert to Greek time
        dt = pytz.UTC.localize(dt).astimezone(pytz.timezone('Europe/Athens'))
    elif dt.tzinfo != pytz.timezone('Europe/Athens'):
        # Convert to Greek time if it's in a different timezone
        dt = dt.astimezone(pytz.timezone('Europe/Athens'))
    
    return dt.strftime(format_str)

class ExportService:
    """Service for exporting transcriptions and model comparisons in PDF/DOCX formats."""
    
    def get_transcriptions_for_export(
        self, 
        user_id: str, 
        start_date: Optional[str] = None,
        end_date: Optional[str] = None,
        models: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        """Get transcription data for export."""
        try:
            query = Transcription.query.filter_by(user_id=user_id)
            
            # Apply date filters
            if start_date:
                start_dt = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
                query = query.filter(Transcription.created_at >= start_dt)
            
            if end_date:
                end_dt = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
                query = query.filter(Transcription.created_at <= end_dt)
            
            # Apply model filters
            if models:
                query = query.filter(Transcription.model_used.in_(models))
            
            transcriptions = query.order_by(Transcription.created_at.desc()).all()
            
            # Convert to dict format
            result = []
            for t in transcriptions:
                result.append({
                    'id': t.id,
                    'filename': t.filename,
                    'model_used': t.model_used,
                    'transcription_text': t.transcription_text,
                    'audio_duration': t.audio_duration,
                    'processing_time': t.processing_time,
                    'accuracy_wer': t.accuracy_wer,
                    'accuracy_cer': t.accuracy_cer,
                    'confidence_score': t.confidence_score,
                    'created_at': t.created_at.isoformat() if t.created_at else None,
                    'language': t.language
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Failed to get transcriptions for export: {str(e)}")
            return []

    def transcriptions_to_pdf(self, transcriptions: List[Dict[str, Any]]) -> bytes:
        """Convert transcriptions to PDF format."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Create custom styles for Greek text
        greek_style = ParagraphStyle(
            'Greek',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            encoding='utf-8'
        )
        
        title_style = ParagraphStyle(
            'GreekTitle',
            parent=styles['Title'],
            fontName='Helvetica-Bold',
            fontSize=16,
            alignment=1,  # Center alignment
            spaceAfter=30
        )
        
        # Title
        title = Paragraph("Αναφορά Μεταγραφών - GreekSTT Comparison Platform", title_style)
        story.append(title)
        story.append(Spacer(1, 20))
        
        # Generate date
        date_text = f"Ημερομηνία Δημιουργίας: {format_greek_datetime()}"
        story.append(Paragraph(date_text, greek_style))
        story.append(Spacer(1, 20))
        
        # Summary
        summary_text = f"Συνολικός Αριθμός Μεταγραφών: {len(transcriptions)}"
        story.append(Paragraph(summary_text, greek_style))
        story.append(Spacer(1, 30))
        
        # Transcriptions
        for i, transcription in enumerate(transcriptions, 1):
            # Transcription header
            header = f"Μεταγραφή {i}: {transcription.get('filename', 'N/A')}"
            story.append(Paragraph(header, styles['Heading2']))
            
            # Transcription details
            details = [
                f"Μοντέλο: {transcription.get('model_used', 'N/A')}",
                f"Διάρκεια Ήχου: {transcription.get('audio_duration', 0):.2f} δευτερόλεπτα",
                f"Χρόνος Επεξεργασίας: {transcription.get('processing_time', 0):.2f} δευτερόλεπτα",
                f"WER: {transcription.get('accuracy_wer', 0):.2f}%",
                f"CER: {transcription.get('accuracy_cer', 0):.2f}%",
                f"Confidence: {transcription.get('confidence_score', 0):.2f}%"
            ]
            
            for detail in details:
                story.append(Paragraph(detail, greek_style))
            
            story.append(Spacer(1, 10))
            
            # Transcription text
            story.append(Paragraph("Κείμενο Μεταγραφής:", styles['Heading3']))
            transcription_text = transcription.get('transcription_text', 'Δεν υπάρχει κείμενο')
            story.append(Paragraph(transcription_text, greek_style))
            
            story.append(Spacer(1, 30))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    def transcriptions_to_docx(self, transcriptions: List[Dict[str, Any]]) -> bytes:
        """Convert transcriptions to DOCX format."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Αναφορά Μεταγραφών - GreekSTT Comparison Platform', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f'Ημερομηνία Δημιουργίας: {format_greek_datetime()}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Summary
        doc.add_paragraph(f'Συνολικός Αριθμός Μεταγραφών: {len(transcriptions)}')
        doc.add_paragraph()
        
        # Transcriptions
        for i, transcription in enumerate(transcriptions, 1):
            # Transcription header
            header = doc.add_heading(f'Μεταγραφή {i}: {transcription.get("filename", "N/A")}', level=2)
            
            # Details table
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            
            details = [
                ('Μοντέλο', transcription.get('model_used', 'N/A')),
                ('Διάρκεια Ήχου', f"{transcription.get('audio_duration', 0):.2f} δευτερόλεπτα"),
                ('Χρόνος Επεξεργασίας', f"{transcription.get('processing_time', 0):.2f} δευτερόλεπτα"),
                ('WER', f"{transcription.get('accuracy_wer', 0):.2f}%"),
                ('CER', f"{transcription.get('accuracy_cer', 0):.2f}%"),
                ('Confidence', f"{transcription.get('confidence_score', 0):.2f}%")
            ]
            
            for row_idx, (label, value) in enumerate(details):
                row = table.rows[row_idx]
                row.cells[0].text = label
                row.cells[1].text = value
            
            # Transcription text
            doc.add_paragraph()
            doc.add_heading('Κείμενο Μεταγραφής:', level=3)
            transcription_text = transcription.get('transcription_text', 'Δεν υπάρχει κείμενο')
            doc.add_paragraph(transcription_text)
            
            doc.add_page_break()
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def get_comparison_for_export(self, comparison_id: str, user_id: str) -> Optional[Dict[str, Any]]:
        """Get comparison data for export."""
        try:
            comparison = ModelComparison.query.filter_by(id=comparison_id, user_id=user_id).first()
            
            if not comparison:
                return None
            
            return {
                'id': comparison.id,
                'whisper_transcription': comparison.whisper_transcription,
                'wav2vec_transcription': comparison.wav2vec_transcription,
                'whisper_wer': comparison.whisper_wer,
                'wav2vec_wer': comparison.wav2vec_wer,
                'whisper_cer': comparison.whisper_cer,
                'wav2vec_cer': comparison.wav2vec_cer,
                'whisper_processing_time': comparison.whisper_processing_time,
                'wav2vec_processing_time': comparison.wav2vec_processing_time,
                'whisper_confidence': comparison.whisper_confidence,
                'wav2vec_confidence': comparison.wav2vec_confidence,
                'audio_duration': comparison.audio_duration,
                'filename': comparison.filename,
                'created_at': comparison.created_at.isoformat() if comparison.created_at else None,
                'language': comparison.language,
                'comparison_insights': comparison.comparison_insights
            }
            
        except Exception as e:
            logger.error(f"Failed to get comparison for export: {str(e)}")
            return None

    def generate_comparison_pdf(self, comparison: Dict[str, Any]) -> bytes:
        """Generate PDF report for model comparison."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Create custom styles for Greek text
        greek_style = ParagraphStyle(
            'Greek',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            encoding='utf-8'
        )
        
        title_style = ParagraphStyle(
            'GreekTitle',
            parent=styles['Title'],
            fontName='Helvetica-Bold',
            fontSize=16,
            alignment=1,  # Center alignment
            spaceAfter=30
        )
        
        # Title
        title = Paragraph("Αναφορά Σύγκρισης Μοντέλων ASR", title_style)
        story.append(title)
        story.append(Spacer(1, 20))
        
        # File info
        story.append(Paragraph(f"Αρχείο: {comparison.get('filename', 'N/A')}", greek_style))
        story.append(Paragraph(f"Ημερομηνία: {format_greek_datetime()}", greek_style))
        story.append(Spacer(1, 20))
        
        # Performance comparison table
        story.append(Paragraph("Σύγκριση Επιδόσεων", styles['Heading2']))
        
        data = [
            ['Μετρική', 'Whisper', 'wav2vec2'],
            ['WER (%)', f"{comparison.get('whisper_wer', 0):.2f}", f"{comparison.get('wav2vec_wer', 0):.2f}"],
            ['CER (%)', f"{comparison.get('whisper_cer', 0):.2f}", f"{comparison.get('wav2vec_cer', 0):.2f}"],
            ['Χρόνος (δευτ.)', f"{comparison.get('whisper_processing_time', 0):.2f}", f"{comparison.get('wav2vec_processing_time', 0):.2f}"],
            ['Confidence (%)', f"{comparison.get('whisper_confidence', 0):.2f}", f"{comparison.get('wav2vec_confidence', 0):.2f}"]
        ]
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 30))
        
        # Transcriptions
        story.append(Paragraph("Μεταγραφές", styles['Heading2']))
        
        story.append(Paragraph("Whisper:", styles['Heading3']))
        whisper_text = comparison.get('whisper_transcription', 'Δεν υπάρχει μεταγραφή')
        story.append(Paragraph(whisper_text, greek_style))
        story.append(Spacer(1, 15))
        
        story.append(Paragraph("wav2vec2:", styles['Heading3']))
        wav2vec_text = comparison.get('wav2vec_transcription', 'Δεν υπάρχει μεταγραφή')
        story.append(Paragraph(wav2vec_text, greek_style))
        story.append(Spacer(1, 15))
        
        # Insights
        if comparison.get('comparison_insights'):
            story.append(Paragraph("Ανάλυση Σύγκρισης:", styles['Heading3']))
            story.append(Paragraph(comparison['comparison_insights'], greek_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    def generate_comparison_docx(self, comparison: Dict[str, Any]) -> bytes:
        """Generate DOCX report for model comparison."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Αναφορά Σύγκρισης Μοντέλων ASR', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # File info
        doc.add_paragraph(f'Αρχείο: {comparison.get("filename", "N/A")}')
        doc.add_paragraph(f'Ημερομηνία: {format_greek_datetime()}')
        doc.add_paragraph()
        
        # Performance comparison
        doc.add_heading('Σύγκριση Επιδόσεων', level=2)
        
        table = doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['Μετρική', 'Whisper', 'wav2vec2']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            table.cell(0, i).paragraphs[0].runs[0].font.bold = True
        
        # Data
        data = [
            ['WER (%)', f"{comparison.get('whisper_wer', 0):.2f}", f"{comparison.get('wav2vec_wer', 0):.2f}"],
            ['CER (%)', f"{comparison.get('whisper_cer', 0):.2f}", f"{comparison.get('wav2vec_cer', 0):.2f}"],
            ['Χρόνος (δευτ.)', f"{comparison.get('whisper_processing_time', 0):.2f}", f"{comparison.get('wav2vec_processing_time', 0):.2f}"],
            ['Confidence (%)', f"{comparison.get('whisper_confidence', 0):.2f}", f"{comparison.get('wav2vec_confidence', 0):.2f}"],
            ['Διάρκεια Ήχου', f"{comparison.get('audio_duration', 0):.2f} δευτ.", f"{comparison.get('audio_duration', 0):.2f} δευτ."]
        ]
        
        for row_idx, row_data in enumerate(data, 1):
            for col_idx, cell_data in enumerate(row_data):
                table.cell(row_idx, col_idx).text = cell_data
        
        doc.add_paragraph()
        
        # Transcriptions
        doc.add_heading('Μεταγραφές', level=2)
        
        doc.add_heading('Whisper:', level=3)
        whisper_text = comparison.get('whisper_transcription', 'Δεν υπάρχει μεταγραφή')
        doc.add_paragraph(whisper_text)
        
        doc.add_heading('wav2vec2:', level=3)
        wav2vec_text = comparison.get('wav2vec_transcription', 'Δεν υπάρχει μεταγραφή')
        doc.add_paragraph(wav2vec_text)
        
        # Insights
        if comparison.get('comparison_insights'):
            doc.add_heading('Ανάλυση Σύγκρισης:', level=3)
            doc.add_paragraph(comparison['comparison_insights'])
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()